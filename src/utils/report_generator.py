import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def format_time(seconds):
    """초 단위 시간을 HH:MM:SS 형태로 변환"""
    if seconds is None:
        return "00:00:00"
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"

def get_audio_duration(audio_path, chunks_data=None):
    """오디오 파일의 총 길이를 초 단위로 획득"""
    try:
        import wave
        # 만약 원본이 wav가 아니더라도 변환된 wav를 찾아서 열기 시도
        check_paths = [audio_path]
        base, ext = os.path.splitext(audio_path)
        check_paths.append(base + "_converted.wav")
        for p in check_paths:
            if os.path.exists(p):
                try:
                    with wave.open(p, "rb") as wf:
                        return wf.getnframes() / float(wf.getframerate())
                except:
                    pass
    except Exception:
        pass
    # Fallback to chunks
    if chunks_data:
        return max([c.get('end', 0) for c in chunks_data] + [0])
    return 0

def get_system_env():
    """시스템 하드웨어 및 구동 환경 정보 조회"""
    try:
        import torch
        gpu_available = torch.cuda.is_available()
        gpu_name = torch.cuda.get_device_name(0) if gpu_available else "N/A"
    except:
        gpu_available = False
        gpu_name = "N/A"
    
    import multiprocessing
    cpu_count = multiprocessing.cpu_count()
    
    return {
        "gpu_status": "🟢 NVIDIA CUDA 가속 활성화" if gpu_available else "🟡 CPU 모드 (CUDA 미활성화)",
        "gpu_name": gpu_name,
        "cpu_threads": f"{cpu_count} Threads"
    }

def set_cell_background(cell, hex_color):
    """표의 특정 셀 배경색 설정"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def apply_text_styling(run, font_name="맑은 고딕", size_pt=10, bold=False, italic=False, color_rgb=None):
    """텍스트 스타일 적용 공통 헬퍼"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def set_report_margins(doc):
    """문서 여백 설정 (상하좌우 1인치)"""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

def add_paragraph_with_spacing(doc, text="", style=None, before=0, after=6, line_spacing=1.15):
    """단락 추가 및 간격 설정"""
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def add_transcription_table(doc, chunks_data, diarization_enabled=True):
    """청크(시간/화자/텍스트) 리스트를 표 형태로 문서에 삽입"""
    if not chunks_data:
        p = add_paragraph_with_spacing(doc, "전사 결과 데이터가 존재하지 않습니다.")
        apply_text_styling(p.add_run(), italic=True)
        return
        
    cols = 3 if diarization_enabled else 2
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    hdr_cells[0].text = '구간 (Timestamp)'
    if diarization_enabled:
        hdr_cells[1].text = '화자 (Speaker)'
        hdr_cells[2].text = '전사 내용 (Transcription)'
    else:
        hdr_cells[1].text = '전사 내용 (Transcription)'
        
    # 헤더 스타일 지정
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, "1F4E78") # Slate Blue
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            apply_text_styling(run, size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
            
    for chunk in chunks_data:
        row_cells = table.add_row().cells
        
        start_t = chunk.get('start', 0)
        end_t = chunk.get('end', 0)
        time_str = f"{format_time(start_t)} - {format_time(end_t)}"
        text = chunk.get('text', '').strip()
        
        row_cells[0].text = time_str
        p_time = row_cells[0].paragraphs[0]
        p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_time.runs:
            apply_text_styling(run, size_pt=9)
            
        if diarization_enabled:
            speaker = chunk.get('speaker', '-')
            row_cells[1].text = str(speaker)
            p_spk = row_cells[1].paragraphs[0]
            p_spk.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_spk.runs:
                apply_text_styling(run, size_pt=9)
                
            row_cells[2].text = text
            p_txt = row_cells[2].paragraphs[0]
            for run in p_txt.runs:
                apply_text_styling(run, size_pt=9)
        else:
            row_cells[1].text = text
            p_txt = row_cells[1].paragraphs[0]
            for run in p_txt.runs:
                apply_text_styling(run, size_pt=9)

def create_stt_report(audio_path, output_dir, stats, chunks_data, chart_image_path=None, task_id="Report", clustering_data=None):
    """
    STT 및 화자 분리 결과에 대한 엔터프라이즈 보고서 형식의 Word 리포트 생성
    """
    doc = Document()
    set_report_margins(doc)
    
    diarization_enabled = stats.get('diarization_enabled', False)
    
    # 1. 문서 제목 및 메타정보
    p_title = add_paragraph_with_spacing(doc, before=24, after=12)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AMEVA STT Enterprise 분석 보고서")
    apply_text_styling(run_title, size_pt=20, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    p_meta = add_paragraph_with_spacing(doc, before=6, after=24)
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"작업 ID: {task_id}\n생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    apply_text_styling(run_meta, size_pt=10, italic=True)
    
    # 2. 분석 개요 및 수행 통계
    p_h1 = add_paragraph_with_spacing(doc, "1. 분석 개요 및 수행 통계", before=18, after=12)
    apply_text_styling(p_h1.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '구분 항목 (Metric)'
    hdr_cells[1].text = '상세 정보 (Value)'
    
    for cell in hdr_cells:
        set_cell_background(cell, "D9E1F2") # Light Blue Shading
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            apply_text_styling(run, size_pt=10, bold=True)
            
    file_size_mb = os.path.getsize(audio_path) / (1024 * 1024) if os.path.exists(audio_path) else 0
    duration_sec = get_audio_duration(audio_path, chunks_data)
    
    stats_data = [
        ("파일명 (File Name)", os.path.basename(audio_path)),
        ("파일 크기 (File Size)", f"{file_size_mb:.2f} MB"),
        ("총 오디오 길이 (Audio Length)", format_time(duration_sec)),
        ("분석 소요 시간 (Processing Time)", f"{stats.get('processing_time_sec', 0):.2f} 초"),
        ("사용 엔진 모델 (Model)", stats.get('model', 'Unknown')),
        ("언어 인식 설정 (Language)", stats.get('language', 'auto')),
        ("화자 분리 활성화 (Diarization)", "활성화 (Vosk)" if diarization_enabled else "비활성화 (STT Only)"),
        ("총 문장(청크) 수 (Total Chunks)", str(len(chunks_data) if chunks_data else 0))
    ]
        
    for k, v in stats_data:
        row_cells = table.add_row().cells
        row_cells[0].text = k
        p_k = row_cells[0].paragraphs[0]
        for r in p_k.runs: apply_text_styling(r, size_pt=10)
        
        row_cells[1].text = str(v)
        p_v = row_cells[1].paragraphs[0]
        for r in p_v.runs: apply_text_styling(r, size_pt=10)
        
    # 시스템 환경 정보 추가
    p_env_title = add_paragraph_with_spacing(doc, "\n* 분석 시스템 환경 정보", before=12, after=6)
    apply_text_styling(p_env_title.add_run(), size_pt=10, bold=True)
    
    env_info = get_system_env()
    p_env_val = add_paragraph_with_spacing(doc, f"- 가속 환경: {env_info['gpu_status']}\n- GPU 모델: {env_info['gpu_name']}\n- CPU 성능: {env_info['cpu_threads']}", before=0, after=12)
    for r in p_env_val.runs: apply_text_styling(r, size_pt=9)
    
    # 3. 화자 분리 시각화 및 군집화 데이터 (활성화 시에만 포함)
    if diarization_enabled and chart_image_path and os.path.exists(chart_image_path):
        doc.add_page_break()
        p_h2 = add_paragraph_with_spacing(doc, "2. 화자 군집화 시각화 분석", before=18, after=12)
        apply_text_styling(p_h2.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
        
        add_paragraph_with_spacing(doc, "다음은 오디오 파일 내 세그먼트별 화자 특징 임베딩(PCA 2차원 축소) 군집 분석 결과입니다. 각 점은 하나의 발화 세그먼트를 나타내며, 동일한 색상은 동일한 화자로 군집화되었음을 나타냅니다.")
        
        # 1:1 종횡비 적용하여 정사각형 차트 삽입 (5인치 x 5인치)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(chart_image_path, width=Inches(5.0), height=Inches(5.0))
        
        # 3-1. 군집화 상세 데이터 표 추가 (또 그 군집화의 데이터 등이 word에 들어가야 함)
        if clustering_data and 'coords' in clustering_data:
            p_h2_sub = add_paragraph_with_spacing(doc, "2.1 화자 군집 데이터 테이블 (Clustering Data)", before=14, after=8)
            apply_text_styling(p_h2_sub.add_run(), size_pt=11, bold=True, color_rgb=RGBColor(0x2E, 0x74, 0xB5))
            
            c_table = doc.add_table(rows=1, cols=4)
            c_table.style = 'Table Grid'
            c_hdr = c_table.rows[0].cells
            c_hdr[0].text = '세그먼트 ID'
            c_hdr[1].text = 'PCA 좌표 (X, Y)'
            c_hdr[2].text = '분류 화자 (Speaker)'
            c_hdr[3].text = '발화 텍스트 요약'
            
            for cell in c_hdr:
                set_cell_background(cell, "D9D9D9")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: apply_text_styling(r, size_pt=9, bold=True)
                
            coords = clustering_data.get('coords', [])
            labels = clustering_data.get('labels', [])
            texts = clustering_data.get('texts', [])
            
            # 너무 많으면 상위 30개만 잘라서 표에 삽입하고 생략 표시
            limit = 40
            display_count = min(len(coords), limit)
            for idx in range(display_count):
                row_cells = c_table.add_row().cells
                row_cells[0].text = f"Seg-{idx+1:03d}"
                row_cells[1].text = f"({coords[idx][0]:.3f}, {coords[idx][1]:.3f})"
                row_cells[2].text = f"Speaker {labels[idx]}"
                
                raw_text = texts[idx] if idx < len(texts) else ""
                # 시간포함된 요약 텍스트인 경우 정제 또는 길이제한
                if "]" in raw_text:
                    raw_text = raw_text.split("]", 1)[1].strip()
                row_cells[3].text = raw_text[:35] + ("..." if len(raw_text) > 35 else "")
                
                for i, cell in enumerate(row_cells):
                    p = cell.paragraphs[0]
                    if i < 3: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: apply_text_styling(r, size_pt=8)
            
            if len(coords) > limit:
                p_note = add_paragraph_with_spacing(doc, f"* 데이터 가독성을 위해 상위 {limit}개 세그먼트 데이터만 표로 출력했습니다. (총 {len(coords)}개)", before=4, after=12)
                apply_text_styling(p_note.add_run(), size_pt=9, italic=True)
                
        doc.add_page_break()
        
    # 4. 전체 전사 결과 테이블
    section_num = "3" if (diarization_enabled and chart_image_path and os.path.exists(chart_image_path)) else "2"
    p_h3 = add_paragraph_with_spacing(doc, f"{section_num}. 타임스탬프별 전사(STT) 결과표", before=18, after=12)
    apply_text_styling(p_h3.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    add_transcription_table(doc, chunks_data, diarization_enabled)
                
    # Save the document
    os.makedirs(output_dir, exist_ok=True)
    report_filename = f"{task_id}_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    report_path = os.path.join(output_dir, report_filename)
    doc.save(report_path)
    
    return report_path


def create_comparison_report(audio_path, output_dir, models_data, task_id="Compare"):
    """
    여러 모델의 결과를 비교하는 엔터프라이즈 보고서 형식의 Word 리포트 생성
    models_data: dict, 키는 모델명, 값은 { 'stats': {}, 'chunks_data': [], 'chart_image_path': str }
    """
    doc = Document()
    set_report_margins(doc)
    
    # Title
    p_title = add_paragraph_with_spacing(doc, before=24, after=12)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AMEVA STT 모델 간 성능 비교 통합 보고서")
    apply_text_styling(run_title, size_pt=20, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    p_meta = add_paragraph_with_spacing(doc, before=6, after=24)
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"대상 파일: {os.path.basename(audio_path)}\n작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    apply_text_styling(run_meta, size_pt=10, italic=True)
    
    # 1. 수행 시스템 환경 정보
    p_h1 = add_paragraph_with_spacing(doc, "1. 평가 환경 구성 정보", before=18, after=12)
    apply_text_styling(p_h1.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    env_info = get_system_env()
    env_table = doc.add_table(rows=1, cols=2)
    env_table.style = 'Table Grid'
    
    ehdr = env_table.rows[0].cells
    ehdr[0].text = '구동 환경 항목'
    ehdr[1].text = '상세 명세'
    for cell in ehdr:
        set_cell_background(cell, "D9E1F2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: apply_text_styling(r, size_pt=10, bold=True)
        
    env_rows = [
        ("컴퓨팅 장치 환경", env_info['gpu_status']),
        ("GPU 상세 정보", env_info['gpu_name']),
        ("CPU 동시 스레드 수", env_info['cpu_threads']),
        ("평가 실행일자", datetime.now().strftime("%Y-%m-%d"))
    ]
    for k, v in env_rows:
        row = env_table.add_row().cells
        row[0].text = k
        row[1].text = v
        for cell in row:
            for r in cell.paragraphs[0].runs: apply_text_styling(r, size_pt=9)
            
    # 2. Executive Summary (성능 비교 요약 표)
    p_h2 = add_paragraph_with_spacing(doc, "\n2. 모델별 주요 성능 평가 요약 (Executive Summary)", before=18, after=12)
    apply_text_styling(p_h2.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    summary_table = doc.add_table(rows=1, cols=6)
    summary_table.style = 'Table Grid'
    hdr = summary_table.rows[0].cells
    hdr[0].text = '평가 모델명'
    hdr[1].text = '오디오 총 길이'
    hdr[2].text = '분석 소요시간'
    hdr[3].text = '인식 속도비 (RTF)'
    hdr[4].text = '생성 문장 수'
    hdr[5].text = '인식 글자 수'
    
    for cell in hdr:
        set_cell_background(cell, "1F4E78")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: apply_text_styling(r, size_pt=9, bold=True, color_rgb=RGBColor(255, 255, 255))
        
    # 오디오 길이 획득
    first_model_chunks = list(models_data.values())[0].get('chunks_data', []) if models_data else []
    duration_sec = get_audio_duration(audio_path, first_model_chunks)
    duration_str = format_time(duration_sec)
    
    for model_name, data in models_data.items():
        stats = data.get('stats', {})
        chunks = data.get('chunks_data', [])
        total_len = sum([len(c.get('text', '')) for c in chunks])
        proc_time = stats.get('processing_time_sec', 0)
        
        # Real-time factor (RTF) 계산: 분석소요시간 / 오디오길이
        rtf = proc_time / duration_sec if duration_sec > 0 else 0
        
        row = summary_table.add_row().cells
        row[0].text = model_name
        row[1].text = duration_str
        row[2].text = f"{proc_time:.2f} 초"
        row[3].text = f"{rtf:.3f}x"
        row[4].text = str(len(chunks))
        row[5].text = f"{total_len} 자"
        
        for i, cell in enumerate(row):
            p = cell.paragraphs[0]
            if i > 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: apply_text_styling(r, size_pt=9)
            
    doc.add_page_break()
    
    # 3. 시간대별 STT 결과 비교표 (같은 시분초에 놓고 한눈에 볼 수 있게)
    p_h3 = add_paragraph_with_spacing(doc, "3. 동일 시간대별 STT 전사 텍스트 정밀 비교", before=18, after=12)
    apply_text_styling(p_h3.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
    
    add_paragraph_with_spacing(doc, "각 모델들이 동일한 타임프레임 안에서 출력한 전사 결과를 대조표 형태로 배치하여, 구간별 인식 오류나 세부 표현 차이를 직접 비교할 수 있습니다.")
    
    valid_models = list(models_data.keys())
    if valid_models:
        # 첫 번째 모델을 기준으로 삼아 타임프레임 기준 설정
        base_model = valid_models[0]
        base_chunks = models_data[base_model].get('chunks_data', [])
        
        comp_table = doc.add_table(rows=1, cols=len(valid_models) + 1)
        comp_table.style = 'Table Grid'
        
        c_hdr = comp_table.rows[0].cells
        c_hdr[0].text = '구간 (Timestamp)'
        set_cell_background(c_hdr[0], "2F5597")
        p = c_hdr[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: apply_text_styling(r, size_pt=9, bold=True, color_rgb=RGBColor(255,255,255))
        
        for idx, m_name in enumerate(valid_models):
            c_hdr[idx + 1].text = m_name
            set_cell_background(c_hdr[idx + 1], "2F5597")
            p = c_hdr[idx + 1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: apply_text_styling(r, size_pt=9, bold=True, color_rgb=RGBColor(255,255,255))
            
        # Keep track of printed/matched chunks to avoid double matching
        printed_chunks = {m: set() for m in valid_models}
        
        for bc in base_chunks:
            t_s, t_e = bc.get("start", 0), bc.get("end", 0)
            time_label = f"[{int(t_s//60):02d}:{int(t_s%60):02d} ~ {int(t_e//60):02d}:{int(t_e%60):02d}]"
            
            row_cells = comp_table.add_row().cells
            row_cells[0].text = time_label
            p_time = row_cells[0].paragraphs[0]
            p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p_time.runs: apply_text_styling(r, size_pt=8)
            
            # Base model text
            row_cells[1].text = bc.get("text", "").strip()
            for r in row_cells[1].paragraphs[0].runs: apply_text_styling(r, size_pt=8)
            
            # Other models text
            for idx, m_name in enumerate(valid_models[1:]):
                m_chunks = models_data[m_name].get('chunks_data', [])
                best_match_idx = None
                best_overlap = 0.0
                
                for mc_idx, mc in enumerate(m_chunks):
                    mc_s, mc_e = mc.get("start", 0), mc.get("end", 0)
                    overlap_start = max(t_s, mc_s)
                    overlap_end = min(t_e, mc_e)
                    overlap = max(0.0, overlap_end - overlap_start)
                    if overlap > best_overlap:
                        best_overlap = overlap
                        best_match_idx = mc_idx
                        
                if best_match_idx is not None:
                    if best_match_idx in printed_chunks[m_name]:
                        row_cells[idx + 2].text = "(이미 위에 포함됨)"
                    else:
                        row_cells[idx + 2].text = m_chunks[best_match_idx].get("text", "").strip()
                        printed_chunks[m_name].add(best_match_idx)
                else:
                    row_cells[idx + 2].text = ""
                for r in row_cells[idx + 2].paragraphs[0].runs: apply_text_styling(r, size_pt=8)
                
    doc.add_page_break()
    
    # 4. 화자 분리 비교 및 개별 상세 (화자 분리가 옵션으로 켜진 경우에만)
    any_diarization = any([data.get('stats', {}).get('diarization_enabled', False) for data in models_data.values()])
    
    if any_diarization:
        p_h4 = add_paragraph_with_spacing(doc, "4. 모델 간 화자 분리 결과 대조 분석", before=18, after=12)
        apply_text_styling(p_h4.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
        
        add_paragraph_with_spacing(doc, "다음 표는 화자 분리가 진행된 각 모델들의 동일 구간 내 화자 세그먼트 매핑 정보입니다.")
        
        # 화자분리 비교 테이블 생성
        dia_comp_table = doc.add_table(rows=1, cols=len(valid_models) + 1)
        dia_comp_table.style = 'Table Grid'
        
        dc_hdr = dia_comp_table.rows[0].cells
        dc_hdr[0].text = '구간 (Timestamp)'
        set_cell_background(dc_hdr[0], "44546A")
        p = dc_hdr[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: apply_text_styling(r, size_pt=9, bold=True, color_rgb=RGBColor(255, 255, 255))
        
        for idx, m_name in enumerate(valid_models):
            dc_hdr[idx + 1].text = f"{m_name} (화자)"
            set_cell_background(dc_hdr[idx + 1], "44546A")
            p = dc_hdr[idx + 1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: apply_text_styling(r, size_pt=9, bold=True, color_rgb=RGBColor(255, 255, 255))
            
        base_model = valid_models[0]
        base_chunks = models_data[base_model].get('chunks_data', [])
        
        # Keep track of printed/matched chunks for diarization to avoid double matching
        printed_dia_chunks = {m: set() for m in valid_models}
        
        for bc in base_chunks:
            t_s, t_e = bc.get("start", 0), bc.get("end", 0)
            time_label = f"[{int(t_s//60):02d}:{int(t_s%60):02d} ~ {int(t_e//60):02d}:{int(t_e%60):02d}]"
            
            row_cells = dia_comp_table.add_row().cells
            row_cells[0].text = time_label
            p_time = row_cells[0].paragraphs[0]
            p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p_time.runs: apply_text_styling(r, size_pt=8)
            
            # Base model diarization
            spk = bc.get("speaker", "Unknown")
            txt = bc.get("text", "").strip()[:15]
            row_cells[1].text = f"[{spk}] {txt}.."
            for r in row_cells[1].paragraphs[0].runs: apply_text_styling(r, size_pt=8)
            
            # Other models diarization
            for idx, m_name in enumerate(valid_models[1:]):
                m_chunks = models_data[m_name].get('chunks_data', [])
                best_match_idx = None
                best_overlap = 0.0
                
                for mc_idx, mc in enumerate(m_chunks):
                    mc_s, mc_e = mc.get("start", 0), mc.get("end", 0)
                    overlap_start = max(t_s, mc_s)
                    overlap_end = min(t_e, mc_e)
                    overlap = max(0.0, overlap_end - overlap_start)
                    if overlap > best_overlap:
                        best_overlap = overlap
                        best_match_idx = mc_idx
                        
                if best_match_idx is not None:
                    if best_match_idx in printed_dia_chunks[m_name]:
                        row_cells[idx + 2].text = "(이미 위에 포함됨)"
                    else:
                        mc = m_chunks[best_match_idx]
                        spk = mc.get("speaker", "Unknown")
                        txt = mc.get("text", "").strip()[:15]
                        row_cells[idx + 2].text = f"[{spk}] {txt}.."
                        printed_dia_chunks[m_name].add(best_match_idx)
                else:
                    row_cells[idx + 2].text = ""
                for r in row_cells[idx + 2].paragraphs[0].runs: apply_text_styling(r, size_pt=8)
                
        doc.add_page_break()
        
        # 각 모델별 화자분리 상세 개별 데이터도 함께 보이기 (정사각형 차트 포함)
        p_h5 = add_paragraph_with_spacing(doc, "5. 모델별 개별 화자 군집화 및 전사 상세", before=18, after=12)
        apply_text_styling(p_h5.add_run(), size_pt=14, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x78))
        
        for model_name, data in models_data.items():
            doc.add_heading(f'■ {model_name} 모델 상세 리포트', level=2)
            stats = data.get('stats', {})
            chunks = data.get('chunks_data', [])
            chart_image_path = data.get('chart_image_path', None)
            
            p_time = add_paragraph_with_spacing(doc, f"분석 완료 소요시간: {stats.get('processing_time_sec', 0):.2f} 초", before=6, after=12)
            apply_text_styling(p_time.runs[0], size_pt=10, italic=True)
            
            if chart_image_path and os.path.exists(chart_image_path):
                p_sub = add_paragraph_with_spacing(doc, f"* {model_name} 화자 군집 임베딩 PCA 시각화 (1:1 정사각형)", before=6, after=6)
                apply_text_styling(p_sub.runs[0], size_pt=9, bold=True)
                
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(chart_image_path, width=Inches(4.5), height=Inches(4.5))
                
            p_tbl_lbl = add_paragraph_with_spacing(doc, f"* {model_name} 상세 전사(STT/화자분리) 결과표", before=12, after=6)
            apply_text_styling(p_tbl_lbl.runs[0], size_pt=9, bold=True)
            add_transcription_table(doc, chunks, diarization_enabled=True)
            doc.add_page_break()
            
    os.makedirs(output_dir, exist_ok=True)
    report_filename = f"Comparison_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    report_path = os.path.join(output_dir, report_filename)
    doc.save(report_path)
    
    return report_path
